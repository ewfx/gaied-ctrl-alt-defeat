import os
import io
import re
import email
from typing import List, Dict, Any, Optional, Tuple
from email.parser import Parser
from email.policy import default
import logging
import base64

# Document processing
import pypdf
import docx
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

class EmailProcessor:
    """
    Service for processing email content and attachments to extract text
    """
    
    def __init__(self, max_attachment_size_mb: int = 10):
        """
        Initialize the email processor
        
        Args:
            max_attachment_size_mb: Maximum attachment size in MB
        """
        self.max_attachment_size = max_attachment_size_mb * 1024 * 1024  # Convert to bytes
    
    def process_email(self, 
                     content: str, 
                     attachments: List[Dict[str, Any]] = None) -> Tuple[str, List[Dict[str, str]]]:
        """
        Process email content and attachments
        
        Args:
            content: Raw email content
            attachments: List of attachment objects with filename, content_type, and content
            
        Returns:
            Tuple of (processed_email_text, processed_attachments_list)
        """
        # Process the email body
        email_text = self.process_email_content(content)
        
        # Process attachments if any
        processed_attachments = []
        if attachments:
            for idx, attachment in enumerate(attachments):
                try:
                    # Check attachment size
                    if len(attachment["content"]) > self.max_attachment_size:
                        logger.warning(f"Attachment too large: {attachment['filename']} ({len(attachment['content'])/1024/1024:.2f} MB)")
                        processed_text = f"[Attachment too large: {attachment['filename']}]"
                    else:
                        processed_text = self.process_attachment(
                            attachment["content"],
                            attachment["filename"],
                            attachment["content_type"]
                        )
                    
                    processed_attachments.append({
                        "index": idx + 1,
                        "filename": attachment["filename"],
                        "content_type": attachment["content_type"],
                        "text": processed_text
                    })
                except Exception as e:
                    logger.error(f"Error processing attachment {attachment['filename']}: {str(e)}")
                    processed_attachments.append({
                        "index": idx + 1,
                        "filename": attachment["filename"],
                        "content_type": attachment["content_type"],
                        "text": f"[Error processing attachment: {str(e)}]"
                    })
        
        return email_text, processed_attachments
    
    def process_email_content(self, content: str) -> str:
        """
        Process raw email content to extract the text body
        
        Args:
            content: Raw email content as string
            
        Returns:
            Extracted text from email body
        """
        # Basic cleaning
        text = self._clean_text(content)
        
        # Try to parse as email if it contains headers
        if "From:" in text or "Subject:" in text or "To:" in text:
            try:
                email_message = Parser(policy=default).parsestr(content)
                body = ""
                
                # Handle multipart emails
                if email_message.is_multipart():
                    for part in email_message.walk():
                        content_type = part.get_content_type()
                        content_disposition = str(part.get("Content-Disposition"))
                        
                        # Skip attachments
                        if "attachment" in content_disposition:
                            continue
                            
                        if content_type == "text/plain":
                            body += part.get_payload(decode=True).decode('utf-8', errors='ignore')
                        elif content_type == "text/html":
                            # Extract text from HTML
                            html_body = part.get_payload(decode=True).decode('utf-8', errors='ignore')
                            body += self._extract_text_from_html(html_body)
                else:
                    # Handle plain text emails
                    content_type = email_message.get_content_type()
                    if content_type == "text/plain":
                        body = email_message.get_payload(decode=True).decode('utf-8', errors='ignore')
                    elif content_type == "text/html":
                        html_body = email_message.get_payload(decode=True).decode('utf-8', errors='ignore')
                        body = self._extract_text_from_html(html_body)
                    else:
                        body = email_message.get_payload()
                        
                if body:
                    return self._clean_text(body)
            except Exception as e:
                logger.error(f"Error parsing email: {e}")
                # Fall back to returning the raw content
        
        return text
    
    def process_attachment(self, 
                          content: bytes, 
                          filename: str, 
                          content_type: str = None) -> str:
        """
        Process attachment file content based on file type
        
        Args:
            content: File content as bytes
            filename: Name of the attachment file
            content_type: MIME type of the content
            
        Returns:
            Extracted text from the attachment
        """
        file_extension = os.path.splitext(filename.lower())[1]
        
        try:
            # Process by file extension
            if file_extension == '.pdf':
                return self._extract_text_from_pdf(content)
            elif file_extension in ['.doc', '.docx']:
                return self._extract_text_from_docx(content)
            elif file_extension == '.txt':
                return content.decode('utf-8', errors='ignore')
            elif file_extension == '.eml':
                # Process as nested email
                return self.process_email_content(content.decode('utf-8', errors='ignore'))
            elif file_extension in ['.htm', '.html']:
                return self._extract_text_from_html(content.decode('utf-8', errors='ignore'))
            elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
                # Image files - return placeholder
                return f"[Image file: {filename}]"
            else:
                # Check content type as fallback
                if content_type and 'pdf' in content_type:
                    return self._extract_text_from_pdf(content)
                elif content_type and ('msword' in content_type or 'document' in content_type):
                    return self._extract_text_from_docx(content)
                elif content_type and 'html' in content_type:
                    return self._extract_text_from_html(content.decode('utf-8', errors='ignore'))
                elif content_type and 'text/plain' in content_type:
                    return content.decode('utf-8', errors='ignore')
                
                return f"[Unsupported file type: {file_extension}]"
        except Exception as e:
            logger.error(f"Error processing attachment {filename}: {str(e)}")
            return f"[Error processing attachment {filename}: {str(e)}]"
    
    def _extract_text_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF file content"""
        text = ""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = pypdf.PdfReader(pdf_file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
                
            return self._clean_text(text)
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            return f"[Error extracting PDF text: {str(e)}]"
    
    def _extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX file content"""
        text = ""
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
                    
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ""
                    for cell in row.cells:
                        row_text += cell.text + " | "
                    text += row_text.strip(" | ") + "\n"
                text += "\n"
                    
            return self._clean_text(text)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            return f"[Error extracting DOCX text: {str(e)}]"
    
    def _extract_text_from_html(self, html_content: str) -> str:
        """Extract text from HTML content using BeautifulSoup"""
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()
                
            # Get text
            text = soup.get_text(separator=' ')
            
            # Clean text
            return self._clean_text(text)
        except Exception as e:
            logger.error(f"Error extracting HTML text: {str(e)}")
            # Fall back to basic regex approach
            text = re.sub(r'<[^>]+>', ' ', html_content)
            return self._clean_text(text)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        if not text:
            return ""
            
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove email forwarding/reply markers
        text = re.sub(r'^>+\s*', '', text, flags=re.MULTILINE)
        
        # Remove common email signature indicators
        text = re.sub(r'--+\s*\n.*$', '', text, flags=re.DOTALL)
        
        # Replace special characters and normalize
        text = text.replace('\r', '\n')
        text = re.sub(r'\t', ' ', text)
        
        return text.strip()